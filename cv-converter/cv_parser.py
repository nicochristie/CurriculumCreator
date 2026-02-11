"""
CV Parser - Extract structure from Word document

This module parses a .docx CV file and extracts:
- Personal information (name, contact details)
- Sections (Experience, Education, Skills, etc.)
- Formatted text (bold, italic)
- Lists and bullet points
- Tables
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class CVParser:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.cv_data = {
            'sections': [],
            'name': '',
            'contact': []
        }

    def parse(self):
        """Main parsing method"""
        # Extract header information from document header
        self._extract_header_info()

        # Extract name/contact from first table if it exists
        self._extract_table_header()

        current_section = None

        for para in self.doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue

            # Determine paragraph type
            if self._is_heading(para):
                # Start a new section
                current_section = {
                    'title': para.text.strip(),
                    'level': self._get_heading_level(para),
                    'content': []
                }
                self.cv_data['sections'].append(current_section)

            elif current_section:
                # Add content to current section
                content_item = self._parse_paragraph(para)
                if content_item:
                    current_section['content'].append(content_item)

            else:
                # Header information (before first section)
                # Could be name, contact info, etc.
                if self._looks_like_name(para):
                    self.cv_data['name'] = para.text.strip()
                elif self._looks_like_contact(para):
                    self.cv_data['contact'].append(self._parse_paragraph(para))

        # Parse tables
        for table in self.doc.tables:
            table_data = self._parse_table(table)
            if current_section:
                current_section['content'].append({
                    'type': 'table',
                    'data': table_data
                })

        return self.cv_data

    def _is_heading(self, para):
        """Detect if paragraph is a heading"""
        text = para.text.strip()

        # Check if paragraph has a heading style
        if para.style.name.startswith('Heading'):
            return True

        # Known section names - whitelist approach
        known_sections = [
            'work experience', 'experience', 'education', 'skills', 'skills summary',
            'projects', 'other', 'certifications', 'awards',
            'publications', 'professional experience', 'technical skills'
        ]

        if text.lower() in known_sections:
            return True

        # "Languages" with tab-separated content is NOT a heading
        if text.lower().startswith('languages') and '\t' in text:
            return False

        # Check for bold, larger font size
        if para.runs:
            first_run = para.runs[0]

            # Exclude text ending with period (likely a sentence, not a heading)
            if text.endswith('.'):
                return False

            # If bold and short, might be heading
            if first_run.bold and len(text) < 30 and len(text.split()) <= 3:
                # But check it's not a job title or similar
                if not any(char in text for char in ['(', ')', '–', '-']) and '@' not in text:
                    return True

            # All caps short text might be a heading
            if text.isupper() and len(text) < 50:
                return True

        return False

    def _get_heading_level(self, para):
        """Determine heading level (1, 2, 3, etc.)"""
        if para.style.name.startswith('Heading'):
            match = re.search(r'Heading (\d+)', para.style.name)
            if match:
                return int(match.group(1))

        # Default to level 2 for detected headings
        return 2

    def _parse_paragraph(self, para):
        """Parse paragraph and extract formatted content"""
        # Check if it's a list item
        is_list_item = self._is_list_item(para)

        # Extract runs with formatting
        runs_data = []
        for run in para.runs:
            if run.text.strip():
                runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                })

        if not runs_data:
            return None

        # Map Word alignment to CSS values
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify'
        }
        css_alignment = alignment_map.get(para.alignment, 'left')

        # Reclassify tech stack lines that were incorrectly identified as list items
        text = para.text.strip()
        if is_list_item and self._is_tech_stack_line(text):
            is_list_item = False

        return {
            'type': 'list_item' if is_list_item else 'paragraph',
            'runs': runs_data,
            'text': text,
            'alignment': css_alignment
        }

    def _is_list_item(self, para):
        """Check if paragraph is a list item"""
        # Check for numbering/bullets in paragraph style
        if para.style.name.startswith('List'):
            return True

        # Check if text starts with bullet characters or numbers
        text = para.text.strip()
        if text and text[0] in ['•', '◦', '▪', '-', '*']:
            return True

        # Check for numbered lists (1., 2., a., etc.)
        if re.match(r'^\d+\.|\w+\.', text):
            return True

        return False

    def _is_tech_stack_line(self, text):
        """Check if text is a technology stack line"""
        # Common tech keywords that indicate this is a tech stack line, not a date/company line
        tech_keywords = [
            'C#', '.NET', 'HTML', 'JavaScript', 'SQL', 'Python', 'Java', 'C++',
            'ASP.NET', 'React', 'Angular', 'Vue', 'Node.js', 'TypeScript',
            'Power Platform', 'SharePoint', 'Azure', 'AWS', 'Firmware',
            'CSS', 'MongoDB', 'PostgreSQL', 'Redis', 'Docker', 'Kubernetes'
        ]

        # Check if text contains multiple tech keywords (high confidence it's a tech line)
        keyword_count = sum(1 for keyword in tech_keywords if keyword in text)

        # If it has 2+ tech keywords, it's likely a tech stack line
        if keyword_count >= 2:
            return True

        # Also check for common tech line patterns
        tech_patterns = [
            r'\b(development|programming|coding)\b',
            r',.*,',  # Multiple commas suggest list of technologies
        ]

        for pattern in tech_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                if keyword_count >= 1:  # At least one tech keyword + pattern
                    return True

        return False

    def _parse_table(self, table):
        """Parse table structure"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    if para.text.strip():
                        cell_content.append(self._parse_paragraph(para))
                row_data.append(cell_content)
            table_data.append(row_data)
        return table_data

    def _looks_like_name(self, para):
        """Heuristic to detect if paragraph contains the person's name"""
        # Usually the first prominent text, possibly bold, centered, or large
        if para.runs and para.runs[0].bold:
            # Short text (likely a name, not a paragraph)
            if len(para.text.split()) <= 4:
                return True

        # Centered text at the beginning
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            if len(para.text.split()) <= 4:
                return True

        return False

    def _looks_like_contact(self, para):
        """Heuristic to detect contact information"""
        text = para.text.lower()
        # Check for email, phone, address keywords
        contact_patterns = [
            r'@',  # email
            r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # phone
            r'\bphone\b', r'\bemail\b', r'\blinkedin\b', r'\bgithub\b',
            r'\baddress\b', r'\bweb\b', r'\bwww\b'
        ]

        for pattern in contact_patterns:
            if re.search(pattern, text):
                return True

        return False

    def _extract_header_info(self):
        """Extract name from document header"""
        for section in self.doc.sections:
            header = section.header
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    # Extract name from "Curriculum Vitae: Name" format
                    if 'curriculum vitae' in text.lower() or 'cv:' in text.lower():
                        # Extract name after colon
                        parts = text.split(':', 1)
                        if len(parts) > 1:
                            name = parts[1].strip()
                            if name and not self.cv_data['name']:
                                self.cv_data['name'] = name
                    elif not self.cv_data['name']:
                        # Just use the header text as name if nothing else found
                        self.cv_data['name'] = text

    def _extract_table_header(self):
        """Extract name and contact info from first table"""
        if not self.doc.tables:
            return

        # Check first table for header info
        first_table = self.doc.tables[0]

        # Typically the header table is small (1-2 rows)
        if len(first_table.rows) <= 2:
            for row in first_table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue

                    # Parse each line in the cell
                    lines = [line.strip() for line in cell_text.split('\n') if line.strip()]

                    for line in lines:
                        # Skip if this line is just the name we already have
                        if self.cv_data['name'] and line == self.cv_data['name']:
                            continue

                        # Check if this line is a name (short, no contact patterns)
                        if not self._looks_like_contact_text(line) and len(line.split()) <= 4:
                            if not self.cv_data['name']:
                                self.cv_data['name'] = line
                        # Otherwise, if it contains contact info, parse it
                        elif self._looks_like_contact_text(line):
                            # Create a simple contact item
                            self.cv_data['contact'].append({
                                'type': 'paragraph',
                                'text': line,
                                'runs': [{'text': line, 'bold': False, 'italic': False, 'underline': False}],
                                'alignment': 'left'
                            })

    def _looks_like_contact_text(self, text):
        """Check if text contains contact information"""
        text_lower = text.lower()
        contact_patterns = [
            r'@',  # email
            r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # phone
            r'\bphone\b', r'\bemail\b', r'\blinkedin\b', r'\bgithub\b',
            r'\baddress\b', r'\bweb\b', r'\bwww\b', r'\bhttp'
        ]

        for pattern in contact_patterns:
            if re.search(pattern, text_lower):
                return True
        return False


def main():
    """Test the parser"""
    parser = CVParser('input/cv.docx')
    cv_data = parser.parse()

    # Print parsed data for verification
    print("=== CV Parsing Results ===\n")
    print(f"Name: {cv_data['name']}")
    print(f"\nContact Info: {len(cv_data['contact'])} items")
    print(f"\nSections: {len(cv_data['sections'])}")

    for section in cv_data['sections']:
        print(f"\n--- {section['title']} (Level {section['level']}) ---")
        print(f"Content items: {len(section['content'])}")

        # Show first few content items
        for i, item in enumerate(section['content'][:3]):
            if item['type'] == 'paragraph':
                print(f"  {i+1}. {item['text'][:60]}...")
            elif item['type'] == 'list_item':
                print(f"  {i+1}. [LIST] {item['text'][:60]}...")
            elif item['type'] == 'table':
                print(f"  {i+1}. [TABLE] {len(item['data'])} rows")


if __name__ == '__main__':
    main()
