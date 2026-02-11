"""Quick script to inspect the CV structure"""
from docx import Document

doc = Document('input/cv.docx')

print("Tables in document:")
print("=" * 60)
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
    # Print first few cells
    for j, row in enumerate(table.rows[:3]):
        for k, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  [{j},{k}]: {cell.text[:60]}")

print("\n\nDocument Header sections:")
print("=" * 60)
for section in doc.sections:
    header = section.header
    for para in header.paragraphs:
        if para.text.strip():
            print(f"Header text: {para.text}")

print("\n\nFirst 10 paragraphs:")
print("=" * 60)
for i, para in enumerate(doc.paragraphs[:10]):
    if para.text.strip():
        print(f"\n[{i}] Style: {para.style.name}")
        print(f"    Text: {para.text[:100]}")
        print(f"    Alignment: {para.alignment}")
        if para.runs:
            print(f"    Bold: {para.runs[0].bold}, Italic: {para.runs[0].italic}")
